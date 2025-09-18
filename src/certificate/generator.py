import os
from pathlib import Path
from docx import Document
import logging
from datetime import datetime
import re
import sys
import time

class CertificateGenerator:
    """Class xử lý tạo giấy khen - hỗ trợ textbox và shapes"""
    
    def __init__(self, template_path, logger=None, config=None):
        self.template_path = Path(template_path)
        self.logger = logger or logging.getLogger(__name__)
        self.config = config
        
        if config:
            self.issued_by = config.get('CERTIFICATE', 'issued_by', fallback='Ban Hướng Dẫn GĐPT')
            self.issued_at = config.get('CERTIFICATE', 'issued_at', fallback='Đà Nẵng')
            self.issued_date = config.get('CERTIFICATE', 'issued_date', fallback='').strip()
            self.no_dharma_name = config.get('CERTIFICATE', 'no_dharma_name', fallback='Không có')
            self.custom_placeholders = {}
            if config.has_section('PLACEHOLDERS'):
                for key, value in config.items('PLACEHOLDERS'):
                    placeholder_key = f'<<{key}>>'
                    self.custom_placeholders[placeholder_key] = value
        else:
            self.issued_by = 'Ban Hướng Dẫn GĐPT'
            self.issued_at = 'Đà Nẵng'
            self.issued_date = ''
            self.no_dharma_name = 'Không có'
            self.custom_placeholders = {}
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"Không tìm thấy template: {template_path}")
    
    def check_template_placeholders(self):
        """Kiểm tra và liệt kê các placeholder trong template - PHIÊN BẢN SIÊU NÂNG CẤP"""
        try:
            doc = Document(str(self.template_path))
            placeholders_found = set()
            debug_info = []
            
            # Hàm tìm placeholder trong text - hỗ trợ cả placeholder bị tách
            def find_placeholders_in_text(text):
                import re
                pattern = r'<<[^>]+>>'
                return re.findall(pattern, text)
            
            def find_placeholders_in_paragraph_runs(paragraph, location=""):
                """Tìm placeholder trong paragraph có thể bị tách thành nhiều run"""
                # Ghép tất cả text từ các run
                full_text = ''.join(run.text for run in paragraph.runs)
                
                # Debug: log chi tiết về runs
                if paragraph.runs and ('<<' in full_text or '>>' in full_text):
                    debug_info.append(f"{location} - Full text: '{full_text}'")
                    for i, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            debug_info.append(f"  Run {i}: '{run.text}'")
                
                return find_placeholders_in_text(full_text)
            
            # Kiểm tra trong paragraphs
            for i, para in enumerate(doc.paragraphs):
                # Kiểm tra cả text thường và text từ runs
                found_normal = find_placeholders_in_text(para.text)
                found_runs = find_placeholders_in_paragraph_runs(para, f"Para {i}")
                placeholders_found.update(found_normal)
                placeholders_found.update(found_runs)
            
            # Kiểm tra trong tables - CHI TIẾT HƠN
            for table_idx, table in enumerate(doc.tables):
                debug_info.append(f"=== TABLE {table_idx + 1} ===")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # Log text trong cell
                        cell_text = cell.text.strip()
                        if cell_text and ('<<' in cell_text or '>>' in cell_text):
                            debug_info.append(f"Cell [{row_idx},{cell_idx}]: '{cell_text}'")
                        
                        for para_idx, para in enumerate(cell.paragraphs):
                            found_normal = find_placeholders_in_text(para.text)
                            found_runs = find_placeholders_in_paragraph_runs(
                                para, f"Table{table_idx+1} Cell[{row_idx},{cell_idx}] Para{para_idx}"
                            )
                            placeholders_found.update(found_normal)
                            placeholders_found.update(found_runs)
            
            # Kiểm tra trong headers/footers
            for section_idx, section in enumerate(doc.sections):
                if section.header:
                    for para_idx, para in enumerate(section.header.paragraphs):
                        found_normal = find_placeholders_in_text(para.text)
                        found_runs = find_placeholders_in_paragraph_runs(
                            para, f"Header{section_idx} Para{para_idx}"
                        )
                        placeholders_found.update(found_normal)
                        placeholders_found.update(found_runs)
                if section.footer:
                    for para_idx, para in enumerate(section.footer.paragraphs):
                        found_normal = find_placeholders_in_text(para.text)
                        found_runs = find_placeholders_in_paragraph_runs(
                            para, f"Footer{section_idx} Para{para_idx}"
                        )
                        placeholders_found.update(found_normal)
                        placeholders_found.update(found_runs)
            
            if self.logger:
                if placeholders_found:
                    self.logger.info("📋 Placeholders tìm thấy trong template:")
                    for p in sorted(placeholders_found):
                        self.logger.info(f"   - {p}")
                else:
                    self.logger.warning("⚠️ Không tìm thấy placeholder nào trong template!")
                    self.logger.info("🔍 Debug chi tiết:")
                    for info in debug_info[:20]:  # Chỉ hiển thị 20 dòng đầu
                        self.logger.info(f"   {info}")
                    if len(debug_info) > 20:
                        self.logger.info(f"   ... và {len(debug_info) - 20} dòng nữa")
            
            return list(placeholders_found)
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"Lỗi kiểm tra template: {e}")
            return []

    def create_certificate(self, ho_ten, phap_danh="", nam_sinh="", don_vi="", output_file=None):
        """Tạo giấy khen - Ưu tiên python-docx như phiên bản cũ"""
        try:
            # Xử lý dữ liệu
            phap_danh_display = phap_danh.strip() if phap_danh.strip() else self.no_dharma_name
            
            if self.issued_date:
                current_date = self.issued_date
            else:
                current_date = datetime.now().strftime("ngày %d tháng %m năm %Y")
            
            if self.logger:
                self.logger.info(f"🔄 Đang xử lý: {ho_ten}")
            
            # Tạo mapping với format ĐÚNG như trong Word template
            replacements = {
                '<<Ho_va_ten>>': ho_ten,
                '<<Phap_danh>>': phap_danh_display,
                '<<Nam_sinh>>': str(nam_sinh) if nam_sinh else '',
                '<<Don_vi>>': don_vi,
                '<<Do>>': self.issued_by,
                '<<Tai>>': self.issued_at,
                '<<Ngay>>': current_date,
            }
            
            # Thêm các custom placeholders
            for key, value in self.custom_placeholders.items():
                replacements[key] = value
            
            if self.logger:
                self.logger.debug("📄 Mapping sẽ sử dụng:")
                for k, v in replacements.items():
                    self.logger.debug(f"  {k} → {v}")
            
            # Kiểm tra template trước khi xử lý
            template_placeholders = self.check_template_placeholders()
            
            # Ưu tiên dùng python-docx như phiên bản cũ đã hoạt động
            success = self._use_python_docx_advanced_v2(replacements, output_file)
            
            # Nếu python-docx thất bại và trên Windows, thử Word COM
            if not success and sys.platform == "win32":
                if self.logger:
                    self.logger.warning("⚠️ python-docx thất bại, thử Word COM...")
                time.sleep(0.5)
                success = self._use_word_com_simple(replacements, output_file)
            
            return success
                
        except Exception as e:
            if self.logger:
                self.logger.error(f"❌ Lỗi tạo giấy khen cho {ho_ten}: {str(e)}")
            return False

    def _use_python_docx_advanced_v2(self, replacements, output_file):
        """Sử dụng python-docx với xử lý run-level cải tiến - PHIÊN BẢN 2"""
        try:
            if self.logger:
                self.logger.info("🔧 Đang sử dụng python-docx (v2)...")
            
            # Đọc template
            doc = Document(str(self.template_path))
            total_replacements = 0
            
            def replace_in_paragraph_v2(paragraph):
                """Thay thế placeholder trong paragraph - THUẬT TOÁN MỚI"""
                if not paragraph.runs:
                    return False
                
                # Bước 1: Ghép tất cả text lại và tìm placeholder
                full_text = ''.join(run.text for run in paragraph.runs)
                
                # Kiểm tra có placeholder không
                has_placeholder = False
                new_text = full_text
                for placeholder, replacement in replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, str(replacement) if replacement else '')
                        has_placeholder = True
                        if self.logger:
                            self.logger.debug(f"   Found & Replaced: {placeholder} → {replacement}")
                
                # Nếu không có placeholder, không làm gì
                if not has_placeholder:
                    return False
                
                # Bước 2: Lưu format của run đầu tiên (hoặc run có text)
                first_run_with_text = None
                for run in paragraph.runs:
                    if run.text.strip():
                        first_run_with_text = run
                        break
                
                if not first_run_with_text:
                    first_run_with_text = paragraph.runs[0]
                
                # Lưu format
                original_format = {
                    'bold': first_run_with_text.bold,
                    'italic': first_run_with_text.italic,
                    'underline': first_run_with_text.underline,
                }
                
                # Lưu font nếu có
                try:
                    original_format['font_name'] = first_run_with_text.font.name
                    original_format['font_size'] = first_run_with_text.font.size
                    original_format['font_color'] = first_run_with_text.font.color.rgb
                except:
                    pass
                
                # Bước 3: Xóa tất cả runs cũ
                paragraph.clear()
                
                # Bước 4: Tạo run mới với text đã thay thế
                new_run = paragraph.add_run(new_text)
                
                # Bước 5: Áp dụng lại format
                try:
                    if original_format.get('bold') is not None:
                        new_run.bold = original_format['bold']
                    if original_format.get('italic') is not None:
                        new_run.italic = original_format['italic']
                    if original_format.get('underline') is not None:
                        new_run.underline = original_format['underline']
                    if original_format.get('font_name'):
                        new_run.font.name = original_format['font_name']
                    if original_format.get('font_size'):
                        new_run.font.size = original_format['font_size']
                    if original_format.get('font_color'):
                        new_run.font.color.rgb = original_format['font_color']
                except Exception as e:
                    if self.logger:
                        self.logger.debug(f"Không thể áp dụng format: {e}")
                
                return True
            
            # Xử lý tất cả paragraphs
            for para in doc.paragraphs:
                if replace_in_paragraph_v2(para):
                    total_replacements += 1
            
            # Xử lý trong tables - CẢI TIẾN ĐẶC BIỆT CHO TABLE
            for table_idx, table in enumerate(doc.tables):
                if self.logger:
                    self.logger.debug(f"Đang xử lý table {table_idx + 1}...")
                
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # Xử lý từng paragraph trong cell
                        for para_idx, para in enumerate(cell.paragraphs):
                            # Debug: hiển thị text trong cell
                            if para.text.strip() and '<<' in para.text:
                                if self.logger:
                                    self.logger.debug(f"  Cell [{row_idx},{cell_idx}] para {para_idx}: '{para.text}'")
                            
                            if replace_in_paragraph_v2(para):
                                total_replacements += 1
                                if self.logger:
                                    self.logger.debug(f"  ✅ Replaced in table cell [{row_idx},{cell_idx}]")
                        
                        # Thêm: Xử lý trực tiếp text trong cell (backup method)
                        try:
                            cell_text = cell.text
                            if any(placeholder in cell_text for placeholder in replacements.keys()):
                                if self.logger:
                                    self.logger.debug(f"  🔄 Trying direct cell text replacement...")
                                # Thử thay thế trực tiếp trong cell text (ít hiệu quả nhưng có thể work)
                                new_cell_text = cell_text
                                for placeholder, replacement in replacements.items():
                                    if placeholder in new_cell_text:
                                        new_cell_text = new_cell_text.replace(placeholder, str(replacement) if replacement else '')
                                
                                # Nếu có thay đổi, clear cell và add lại
                                if new_cell_text != cell_text:
                                    # Clear tất cả paragraphs trong cell
                                    for para in cell.paragraphs[::-1]:  # Reverse để tránh index issues
                                        if len(cell.paragraphs) > 1:
                                            cell._element.remove(para._element)
                                    
                                    # Add text mới vào paragraph đầu tiên
                                    if cell.paragraphs:
                                        cell.paragraphs[0].clear()
                                        cell.paragraphs[0].add_run(new_cell_text)
                                    else:
                                        # Tạo paragraph mới nếu cần
                                        new_para = cell.add_paragraph()
                                        new_para.add_run(new_cell_text)
                                    
                                    total_replacements += 1
                                    if self.logger:
                                        self.logger.debug(f"  ✅ Direct cell replacement successful")
                        except Exception as e:
                            if self.logger:
                                self.logger.debug(f"  ⚠️ Direct cell replacement failed: {e}")
            
            # Xử lý headers và footers
            for section in doc.sections:
                # Header
                if section.header:
                    for para in section.header.paragraphs:
                        if replace_in_paragraph_v2(para):
                            total_replacements += 1
                
                # Footer
                if section.footer:
                    for para in section.footer.paragraphs:
                        if replace_in_paragraph_v2(para):
                            total_replacements += 1
            
            if self.logger:
                self.logger.info(f"📊 python-docx v2 - Tổng thay thế: {total_replacements} vị trí")
            
            # Lưu file
            if output_file:
                # Đảm bảo thư mục tồn tại
                output_file = Path(output_file)
                output_file.parent.mkdir(parents=True, exist_ok=True)
                
                # Lưu document
                doc.save(str(output_file))
                
                if self.logger:
                    self.logger.info(f"✅ Tạo thành công bằng python-docx v2: {output_file.name}")
                
                return True if total_replacements > 0 else False
            
            return False
                
        except Exception as e:
            if self.logger:
                self.logger.error(f"❌ Lỗi python-docx v2: {e}")
            return False

    def _use_python_docx_advanced(self, replacements, output_file):
        """Sử dụng python-docx với xử lý run-level cải tiến - PHIÊN BẢN CŨ (backup)"""
        try:
            if self.logger:
                self.logger.info("🔧 Đang sử dụng python-docx...")
            
            # Đọc template
            doc = Document(str(self.template_path))
            total_replacements = 0
            
            # Hàm thay thế text trong paragraph - xử lý cả trường hợp placeholder bị tách
            def replace_in_paragraph(paragraph):
                """Thay thế placeholder trong paragraph, xử lý cả khi bị tách thành nhiều run"""
                if not paragraph.runs:
                    return False
                    
                # Lưu format của từng run
                run_formats = []
                for run in paragraph.runs:
                    run_formats.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                    })
                
                # Ghép tất cả text lại
                full_text = ''.join(run.text for run in paragraph.runs)
                
                # Kiểm tra và thay thế
                new_text = full_text
                replaced = False
                for placeholder, replacement in replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, str(replacement) if replacement else '')
                        replaced = True
                        if self.logger:
                            self.logger.debug(f"   Replaced: {placeholder} → {replacement}")
                
                # Nếu có thay đổi, cập nhật paragraph
                if replaced:
                    # Giữ format của run đầu tiên
                    first_format = run_formats[0] if run_formats else None
                    
                    # Xóa tất cả runs cũ
                    paragraph.clear()
                    
                    # Tạo run mới với text đã thay thế
                    new_run = paragraph.add_run(new_text)
                    
                    # Áp dụng lại format
                    if first_format:
                        if first_format['bold'] is not None:
                            new_run.bold = first_format['bold']
                        if first_format['italic'] is not None:
                            new_run.italic = first_format['italic']
                        if first_format['underline'] is not None:
                            new_run.underline = first_format['underline']
                        if first_format['font_name']:
                            new_run.font.name = first_format['font_name']
                        if first_format['font_size']:
                            new_run.font.size = first_format['font_size']
                
                return replaced
            
            # Xử lý tất cả paragraphs
            for para in doc.paragraphs:
                if replace_in_paragraph(para):
                    total_replacements += 1
            
            # Xử lý trong tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if replace_in_paragraph(para):
                                total_replacements += 1
            
            # Xử lý headers và footers
            for section in doc.sections:
                # Header
                if section.header:
                    for para in section.header.paragraphs:
                        if replace_in_paragraph(para):
                            total_replacements += 1
                
                # Footer
                if section.footer:
                    for para in section.footer.paragraphs:
                        if replace_in_paragraph(para):
                            total_replacements += 1
            
            if self.logger:
                self.logger.info(f"📊 python-docx - Tổng thay thế: {total_replacements} vị trí")
            
            # Lưu file
            if output_file:
                # Đảm bảo thư mục tồn tại
                output_file = Path(output_file)
                output_file.parent.mkdir(parents=True, exist_ok=True)
                
                # Lưu document
                doc.save(str(output_file))
                
                if self.logger:
                    self.logger.info(f"✅ Tạo thành công bằng python-docx: {output_file.name}")
                
                return True if total_replacements > 0 else False
            
            return False
                
        except Exception as e:
            if self.logger:
                self.logger.error(f"❌ Lỗi python-docx: {e}")
            return False

    def _use_word_com_simple(self, replacements, output_file):
        """Sử dụng Word COM với DEBUG MODE - hiển thị Word để xem chuyện gì xảy ra"""
        try:
            import win32com.client
        except ImportError:
            if self.logger:
                self.logger.warning("⚠️ Không có win32com")
            return False

        word = None
        doc = None
        try:
            if self.logger:
                self.logger.info("🔧 Đang sử dụng Word COM DEBUG MODE...")

            # Khởi tạo Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = True  # HIỂN THỊ Word để debug
            word.DisplayAlerts = 0
            
            # Mở template gốc trực tiếp (không copy)
            doc = word.Documents.Open(
                str(self.template_path.resolve()),
                ReadOnly=False,
                AddToRecentFiles=False,
                Visible=True
            )
            
            if self.logger:
                self.logger.info(f"📄 Document opened: {doc.Name}")
                self.logger.info("🔍 Word đang hiển thị - bạn có thể xem document!")
                
            # Đợi user xác nhận
            input("👁️ Nhấn Enter sau khi bạn đã xem document trong Word...")
            
            total_replacements = 0
            
            # Thử replace từng placeholder một và kiểm tra ngay
            for placeholder, replacement in replacements.items():
                try:
                    if self.logger:
                        self.logger.info(f"🔄 Đang xử lý: {placeholder}")
                    
                    # Sử dụng Selection để thay thế (giống như user làm thủ công)
                    selection = word.Selection
                    find = selection.Find
                    
                    # Reset về đầu document
                    selection.HomeKey(6)  # wdStory = 6
                    
                    # Cấu hình find
                    find.ClearFormatting()
                    find.Replacement.ClearFormatting()
                    find.Text = placeholder
                    find.Replacement.Text = str(replacement) if replacement else ''
                    find.Forward = True
                    find.Wrap = 1  # wdFindContinue = 1
                    find.Format = False
                    find.MatchCase = False
                    find.MatchWholeWord = False
                    
                    # Thực hiện replace all
                    replaced = find.Execute(Replace=2)  # wdReplaceAll = 2
                    
                    if replaced:
                        total_replacements += 1
                        if self.logger:
                            self.logger.info(f"   ✅ {placeholder} → {replacement}")
                        
                        # Đợi để user có thể thấy thay đổi
                        input(f"👀 Đã thay thế {placeholder}. Nhấn Enter để tiếp tục...")
                    else:
                        if self.logger:
                            self.logger.warning(f"   ❌ Không tìm thấy: {placeholder}")
                
                except Exception as e:
                    if self.logger:
                        self.logger.error(f"   ⚠️ Lỗi xử lý {placeholder}: {e}")
            
            if self.logger:
                self.logger.info(f"📊 Word COM DEBUG - Tổng thay thế: {total_replacements} vị trí")
            
            # Đợi user kiểm tra kết quả cuối cùng
            input("🔍 Kiểm tra document cuối cùng trong Word. Nhấn Enter để lưu...")
            
            # Lưu file
            if output_file:
                output_file = Path(output_file)
                output_file.parent.mkdir(parents=True, exist_ok=True)
                doc.SaveAs2(str(output_file.resolve()))
                if self.logger:
                    self.logger.info(f"💾 Đã lưu file: {output_file.name}")
            
            return True if total_replacements > 0 else False
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"❌ Lỗi Word COM DEBUG: {e}")
            return False
        finally:
            # Đóng document và Word
            try:
                if doc:
                    doc.Close(False)
            except:
                pass
            try:
                if word:
                    word.Quit()
            except:
                pass

    def batch_create(self, data_list, output_folder):
        """Tạo nhiều giấy khen cùng lúc"""
        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)
        
        success_count = 0
        failed_list = []
        
        for idx, data in enumerate(data_list, 1):
            try:
                ho_ten = data.get('ho_ten', '')
                phap_danh = data.get('phap_danh', '')
                nam_sinh = data.get('nam_sinh', '')
                don_vi = data.get('don_vi', '')
                
                safe_name = ho_ten.replace(' ', '_').replace('/', '_').replace('\\', '_')
                output_file = output_folder / f"{idx:03d}_{safe_name}.docx"
                
                if self.create_certificate(ho_ten, phap_danh, nam_sinh, don_vi, output_file):
                    success_count += 1
                else:
                    failed_list.append(ho_ten)
                    
            except Exception as e:
                failed_list.append(data.get('ho_ten', f'Record {idx}'))
                if self.logger:
                    self.logger.error(f"❌ Lỗi xử lý {data.get('ho_ten', '')}: {str(e)}")
        
        return success_count, failed_list